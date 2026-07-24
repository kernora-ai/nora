#!/usr/bin/env python3
# Kernora — AI Work Intelligence
# Elastic License 2.0 — commercial use requires agreement with kernora.ai
# https://github.com/kernora-ai/nora/blob/main/LICENSE
"""
Kernora — Document Intelligence Ingestion
Scans project directories for AI-generated / work documents (.md, .pdf, .docx, .txt),
extracts text, runs LLM intelligence extraction, and stores in the `documents` DB table.

Usage:
    python3 doc_ingest.py               # analyze new/changed docs only
    python3 doc_ingest.py --force       # re-analyze everything
    python3 doc_ingest.py --paths ~/code ~/Documents
"""
from __future__ import annotations

import hashlib
import json
import os
import re
import sqlite3
import sys
import time
from datetime import datetime
from pathlib import Path
from typing import Optional

# ── Config loader (mirrors dashboard.py approach) ────────────────────────────

def _load_cfg() -> dict:
    cfg_path = Path.home() / ".kernora" / "config.toml"
    if not cfg_path.exists():
        return {}
    raw = cfg_path.read_text()
    if sys.version_info >= (3, 11):
        import tomllib
        return tomllib.loads(raw)
    try:
        import tomllib
        return tomllib.loads(raw)
    except ImportError:
        try:
            import tomli as tomllib  # type: ignore
            return tomllib.loads(raw)
        except ImportError:
            pass
    # Minimal hand-rolled TOML parser for [section] key = "val"
    cfg: dict = {}
    section = None
    for line in raw.splitlines():
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        m = re.match(r'^\[(\w+)\]$', line)
        if m:
            section = m.group(1)
            cfg.setdefault(section, {})
            continue
        m = re.match(r'^(\w+)\s*=\s*"([^"]*)"$', line)
        if m:
            k, v = m.group(1), m.group(2)
            if section:
                cfg[section][k] = v
            else:
                cfg[k] = v
    return cfg


# ── DB setup ─────────────────────────────────────────────────────────────────

def ensure_documents_table(conn: sqlite3.Connection):
    conn.execute("""
        CREATE TABLE IF NOT EXISTS documents (
            id           TEXT PRIMARY KEY,
            path         TEXT UNIQUE NOT NULL,
            filename     TEXT,
            project      TEXT,
            file_type    TEXT,
            size_bytes   INTEGER,
            content_hash TEXT,
            raw_text     TEXT,
            information_map TEXT,
            doc_type     TEXT,
            summary      TEXT,
            themes       TEXT,
            analyzed_at  TEXT,
            created_at   TEXT DEFAULT (datetime('now'))
        )
    """)
    conn.commit()


# ── Text extraction ───────────────────────────────────────────────────────────

def _extract_md(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _extract_pdf(path: Path) -> str:
    # Try pdfminer first (best quality)
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        text = pdfminer_extract(str(path))
        if text and len(text.strip()) > 50:
            return text
    except Exception:
        pass
    # Fallback: PyPDF2
    try:
        import PyPDF2
        parts = []
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    parts.append(t)
        return "\n".join(parts)
    except Exception as e:
        return f"[PDF extraction failed: {e}]"


def _extract_docx(path: Path) -> str:
    try:
        import docx
        doc = docx.Document(str(path))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        return f"[DOCX extraction failed: {e}]"


def extract_text(path: Path) -> Optional[str]:
    ext = path.suffix.lower()
    try:
        if ext in (".md", ".markdown", ".txt", ".text", ".rst"):
            return _extract_md(path)
        elif ext == ".pdf":
            return _extract_pdf(path)
        elif ext == ".docx":
            return _extract_docx(path)
    except Exception as e:
        print(f"  [!] Extract error {path.name}: {e}")
    return None


# ── Scanner ───────────────────────────────────────────────────────────────────

SKIP_DIRS = {
    ".git", ".svn", "node_modules", "__pycache__", ".venv", "venv", "env",
    ".env", "dist", "build", ".next", ".nuxt", "coverage", ".pytest_cache",
    ".mypy_cache", ".tox", "target", "vendor", "Pods", ".build",
    ".kernora", ".claude", "Library", "Applications",
}

SKIP_FILENAMES = {
    "CHANGELOG.md", "CHANGELOG.txt", "LICENSE", "LICENSE.md", "LICENSE.txt",
    "CONTRIBUTING.md", "CODE_OF_CONDUCT.md", "SECURITY.md",
    ".gitignore", ".npmignore", "requirements.txt", "package-lock.json",
    "yarn.lock", "Pipfile.lock", "poetry.lock", "Cargo.lock",
}

INCLUDE_EXTS = {".md", ".markdown", ".txt", ".pdf", ".docx", ".rst"}
MIN_SIZE = 300       # bytes — skip tiny files
MAX_SIZE = 10_000_000  # 10 MB — skip huge files


def should_skip(path: Path) -> bool:
    if path.name in SKIP_FILENAMES:
        return True
    if path.name.startswith("."):
        return True
    try:
        sz = path.stat().st_size
        if sz < MIN_SIZE or sz > MAX_SIZE:
            return True
    except Exception:
        return True
    return False


def scan_dirs(roots: list[str], max_files: int = 2000) -> list[Path]:
    found: list[Path] = []
    for root_str in roots:
        root = Path(root_str).expanduser().resolve()
        if not root.exists():
            print(f"  [skip] {root} — does not exist")
            continue
        count_before = len(found)
        for dirpath_str, dirnames, filenames in os.walk(root):
            dirpath = Path(dirpath_str)
            dirnames[:] = [
                d for d in dirnames
                if d not in SKIP_DIRS and not d.startswith(".")
            ]
            for fname in filenames:
                fpath = dirpath / fname
                if fpath.suffix.lower() in INCLUDE_EXTS and not should_skip(fpath):
                    found.append(fpath)
                    if len(found) >= max_files:
                        print(f"  [warn] Hit {max_files} file limit — stopping scan")
                        return found
        print(f"  {root}: {len(found) - count_before} documents found")
    return found


# ── LLM analysis ─────────────────────────────────────────────────────────────

_ANALYSIS_PROMPT = """\
You are extracting intelligence from a work document for a founder's knowledge base.
Return ONLY valid JSON (no markdown, no explanation). Use this exact schema:

{
  "document_type": "spec|report|analysis|research|plan|notes|design|legal|financial|other",
  "summary": "2-3 sentence summary",
  "domain_classification": {
    "primary": "engineering|product|business|design|research|marketing|legal|finance|operations|personal",
    "secondary": []
  },
  "key_discussions": [
    {"topic": "...", "conclusion": "...", "status": "decided|open|pending"}
  ],
  "learnings": [
    {"insight": "...", "domain": "..."}
  ],
  "product_intelligence": {
    "tenets": [{"tenet": "...", "product": "..."}],
    "user_insights": [],
    "feature_decisions": [],
    "positioning": ""
  },
  "process_intelligence": {
    "processes_defined": [{"name": "...", "steps": "..."}],
    "tool_skill_candidates": [{"name": "...", "rationale": "...", "complexity": "simple|moderate|complex"}],
    "workflow_patterns": []
  },
  "engineering_intelligence": {
    "system_design": [],
    "tech_debt": [],
    "performance_targets": [],
    "infrastructure": []
  },
  "entities": {
    "people": [],
    "products": [],
    "technologies": [],
    "organizations": [],
    "metrics": []
  },
  "open_questions": [],
  "themes": []
}

Only populate fields that have real content from the document. Return {} for empty sections.\
"""


def _resolve_model(cfg: dict) -> tuple[str, str]:
    """Return (model_string, api_key_env_var)."""
    provider = (cfg.get("model") or {}).get("provider", "anthropic")
    if provider == "grok":
        model = (cfg.get("grok") or {}).get("model", "xai/grok-4-1-fast-reasoning")
        return model, "XAI_API_KEY"
    elif provider == "gemini":
        model = (cfg.get("gemini") or {}).get("model", "gemini/gemini-2.0-flash")
        return model, "GEMINI_API_KEY"
    else:
        model = (cfg.get("anthropic") or {}).get("model", "claude-haiku-4-5-20251001")
        return model, "ANTHROPIC_API_KEY"


def analyze_document(text: str, filename: str, cfg: dict) -> dict:
    try:
        from litellm import completion
    except ImportError:
        print("  [!] litellm not installed — skipping LLM analysis")
        return {}

    model, _ = _resolve_model(cfg)

    # Truncate to ~40k chars (~10k tokens) so we stay in safe context range
    max_chars = 40_000
    body = text[:max_chars]
    if len(text) > max_chars:
        body += f"\n\n[... truncated at {max_chars} chars ...]"

    # 086-B: Wrap document in <document> XML tags with <source> and <document_content> subtags.
    # From Anthropic best practices: "Wrap each document in <document> tags with <source> and
    # <document_content> subtags" for cleaner document boundary parsing.
    user_content = f"""<documents>
  <document index="1">
    <source>{filename} ({cfg.get('file_type', 'document')})</source>
    <document_content>
{body}
    </document_content>
  </document>
</documents>

{_ANALYSIS_PROMPT}"""
    messages = [{"role": "user", "content": user_content}]

    for attempt in range(3):
        try:
            resp = completion(model=model, messages=messages, timeout=120, temperature=0.1)
            raw = (resp.choices[0].message.content or "").strip()
            # Strip markdown code fences if present
            raw = re.sub(r'^```(?:json)?\s*', '', raw)
            raw = re.sub(r'\s*```$', '', raw)
            m = re.search(r'\{.*\}', raw, re.DOTALL)
            if m:
                return json.loads(m.group())
            return {}
        except Exception as e:
            if attempt < 2:
                time.sleep(2 ** attempt)
            else:
                print(f"  [!] LLM failed for {filename}: {e}")
    return {}


# ── Project guesser ───────────────────────────────────────────────────────────

_CODE_ROOTS = {"code", "projects", "work", "dev", "src", "repos", "workspace"}


def _guess_project(path: Path) -> str:
    parts = path.parts
    for i, part in enumerate(parts):
        if part.lower() in _CODE_ROOTS and i + 1 < len(parts):
            return parts[i + 1]
    # Fall back: grandparent dir name
    return path.parent.parent.name if len(parts) >= 3 else path.parent.name


# ── Main ingest ───────────────────────────────────────────────────────────────

def ingest_documents(
    project_paths: list[str],
    db_path: str,
    cfg: dict,
    force: bool = False,
    max_files: int = 2000,
) -> dict:
    """
    Scan project_paths → extract text → LLM analysis → store in documents table.
    Returns summary: {scanned, analyzed, errors, skipped}
    """
    conn = sqlite3.connect(db_path)
    ensure_documents_table(conn)

    all_files = scan_dirs(project_paths, max_files=max_files)
    print(f"\nTotal candidate documents: {len(all_files)}")

    # Filter already-done
    if not force:
        existing = set(
            row[0] for row in
            conn.execute(
                "SELECT path FROM documents WHERE information_map IS NOT NULL AND information_map != '{}'"
            ).fetchall()
        )
        pending = [f for f in all_files if str(f) not in existing]
        print(f"Pending: {len(pending)}  (already done: {len(all_files) - len(pending)})")
    else:
        pending = list(all_files)
        print(f"Force mode: re-analyzing all {len(pending)}")

    done = errors = skipped = 0
    start = time.time()

    for i, path in enumerate(pending):
        print(f"[{i+1}/{len(pending)}] {path.name} ...", end=" ", flush=True)
        t0 = time.time()
        try:
            text = extract_text(path)
            if not text or len(text.strip()) < 100:
                print("skip (too short)")
                skipped += 1
                continue

            content_hash = hashlib.md5(text.encode("utf-8", errors="ignore")).hexdigest()

            # Skip if unchanged
            row = conn.execute(
                "SELECT content_hash FROM documents WHERE path = ?", (str(path),)
            ).fetchone()
            if row and row[0] == content_hash and not force:
                print("skip (unchanged)")
                skipped += 1
                continue

            imap = analyze_document(text, path.name, cfg)
            doc_id = hashlib.md5(str(path).encode()).hexdigest()
            project = _guess_project(path)
            doc_type = imap.get("document_type", "")
            summary = imap.get("summary", "")
            themes = json.dumps(imap.get("themes", []))

            conn.execute("""
                INSERT INTO documents
                    (id, path, filename, project, file_type, size_bytes,
                     content_hash, raw_text, information_map, doc_type, summary,
                     themes, analyzed_at)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,datetime('now'))
                ON CONFLICT(path) DO UPDATE SET
                    content_hash    = excluded.content_hash,
                    raw_text        = excluded.raw_text,
                    information_map = excluded.information_map,
                    doc_type        = excluded.doc_type,
                    summary         = excluded.summary,
                    themes          = excluded.themes,
                    analyzed_at     = excluded.analyzed_at
            """, (
                doc_id, str(path), path.name, project,
                path.suffix.lower(), path.stat().st_size,
                content_hash, text[:80_000],
                json.dumps(imap), doc_type, summary, themes,
            ))
            conn.commit()
            done += 1
            elapsed = int(time.time() - t0)
            rate = done / max((time.time() - start) / 60, 0.01)
            remaining = (len(pending) - i - 1) / max(rate, 0.01)
            print(f"done ({elapsed}s) | {done} done, ~{remaining:.0f}min left")

        except Exception as e:
            print(f"ERROR: {e}")
            errors += 1

    conn.close()
    result = {"scanned": len(all_files), "analyzed": done, "errors": errors, "skipped": skipped}
    print(f"\n{'='*50}")
    print(f"Document ingestion complete: {result}")
    return result


# ── Pull intelligence for broadcast/CLAUDE.md ─────────────────────────────────

def pull_document_intelligence(conn: sqlite3.Connection) -> dict:
    """Pull structured intelligence from analyzed documents."""
    rows = conn.execute("""
        SELECT filename, project, file_type, information_map, summary
        FROM documents
        WHERE information_map IS NOT NULL AND information_map NOT IN ('{}', '')
        ORDER BY analyzed_at DESC
        LIMIT 300
    """).fetchall()

    tenets, learnings, processes, decisions, open_qs = [], [], [], [], []

    for row in rows:
        filename, project, ftype, imap_raw, summary = row
        try:
            imap = json.loads(imap_raw) if imap_raw else {}
        except Exception:
            continue

        pi = imap.get("product_intelligence") or {}
        for t in pi.get("tenets") or []:
            text = t.get("tenet") or (t if isinstance(t, str) else "")
            if text and len(text) > 15:
                tenets.append({"text": text, "project": project or t.get("product",""), "source": filename})

        for l in imap.get("learnings") or []:
            text = l.get("insight") or (l if isinstance(l, str) else "")
            if text and len(text) > 15:
                learnings.append({"text": text, "domain": l.get("domain",""), "source": filename})

        pri = imap.get("process_intelligence") or {}
        for p in pri.get("processes_defined") or []:
            name = p.get("name") or (p if isinstance(p, str) else "")
            if name and len(name) > 5:
                processes.append({"name": name, "steps": p.get("steps",""), "source": filename})

        for kd in imap.get("key_discussions") or []:
            text = kd.get("topic") or kd.get("conclusion") or (kd if isinstance(kd, str) else "")
            if text and len(text) > 10:
                decisions.append({"text": text, "status": kd.get("status",""), "source": filename})

        for q in imap.get("open_questions") or []:
            qs = q if isinstance(q, str) else q.get("question","")
            if qs and len(qs) > 10:
                open_qs.append({"text": qs, "source": filename})

    return {
        "tenets": tenets,
        "learnings": learnings,
        "processes": processes,
        "decisions": decisions,
        "open_questions": open_qs,
        "total_docs": len(rows),
    }


# ── CLI ───────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    cfg = _load_cfg()
    db_path = str(Path.home() / ".kernora" / "echo.db")
    force = "--force" in sys.argv

    # Parse --paths flag
    if "--paths" in sys.argv:
        idx = sys.argv.index("--paths")
        project_paths = sys.argv[idx + 1:]
    else:
        # Default: scan ~/code and ~/Documents
        project_paths = [
            os.path.expanduser("~/code"),
            os.path.expanduser("~/Documents"),
        ]

    print(f"Kernora Document Ingest — {datetime.now().strftime('%H:%M:%S')}")
    print(f"DB: {db_path}")
    print(f"Paths: {project_paths}")
    print(f"Force: {force}")
    print()

    result = ingest_documents(
        project_paths=project_paths,
        db_path=db_path,
        cfg=cfg,
        force=force,
    )
    print(f"\nFinal result: {result}")
